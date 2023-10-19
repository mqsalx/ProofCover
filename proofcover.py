# flake8: noqa
# type: ignore

# from docx import Document
# from docx.shared import Pt
# from docx.shared import RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# # Cria um novo documento
# doc = Document()

# # Adiciona um título ao documento com fonte Arial, tamanho 48, centralizado e cor preta (sem negrito)
# titulo = doc.add_paragraph('Título do Documento')
# titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto
# run = titulo.runs[0]
# run.font.name = 'Arial'  # Define a fonte como Arial
# run.font.size = Pt(48)  # Define o tamanho da fonte como 48 pontos
# run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta

# # Adiciona mais espaço após o título (três quebras de linha)
# doc.add_paragraph('\n\n\n')

# # Adiciona uma tabela ao documento e a centraliza
# table = doc.add_table(rows=10, cols=3)
# table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza a tabela

# # Ajusta a largura da primeira coluna para 1.75 polegadas (por exemplo) e as outras colunas para 1 polegada
# table.columns[0].width = 1.75
# table.columns[1].width = 1
# table.columns[2].width = 1

# # Dados para preencher a tabela
# dados = [
#     ("BIMESTRE", "-", "4º"),
#     ("TIPO DE PROVA", "-", "P1 OBJ"),
#     ("SÉRIE", "-", "1A"),
#     ("DISCIPLINA", "-", "MATEMÁTICA"),
#     ("PROFESSOR", "-", "THIAGO"),
#     ("Nº DE PROVAS", "-", "33"),
#     ("Nº DE 2ª CH", "-", "1"),
#     ("DATA APLICAÇÃO", "-", "09/10/23"),
#     ("DATA DIGITALIZAÇÃO", "-", "11/10/23"),
#     ("ENTREGUE POR", "-", "ESTEVAN"),
#     ("RECEBIDO POR", "-", "ALEX"),
#     ("APLICADOR", "-", "EDUARDA")
# ]

# # Define a fonte, o tamanho e converte o texto para maiúsculas na tabela
# for row_idx, row in enumerate(table.rows):
#     for col_idx, cell in enumerate(row.cells):
#         texto = dados[row_idx][col_idx].upper()  # Converte o texto para maiúsculas
#         cell.text = texto  # Preenche a célula com os dados correspondentes
#         cell.paragraphs[0].runs[0].font.name = 'Arial'  # Define a fonte como Arial
#         cell.paragraphs[0].runs[0].font.size = Pt(22)  # Define o tamanho da fonte para 22 pontos
#         cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta
        
#         # Alinha o texto à esquerda na primeira coluna
#         if col_idx == 0:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#             cell.margin_left = Pt(0)  # Remove qualquer recuo na margem esquerda
#         # Ajusta o recuo para a direita na coluna do meio
#         elif col_idx == 1:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             cell.margin_left = Pt(20)  # Ajusta o recuo para a direita
#         else:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto nas outras colunas

# # Salva o documento
# doc.save('tabela_com_titulo_e_texto_maiusculo_e_alinhamento_sem_recuos.docx')

# print('Documento criado com sucesso!')

# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# # Cria um novo documento
# doc = Document()

# # Adiciona um título ao documento com fonte Arial, tamanho 48, centralizado e cor preta (sem negrito)
# titulo = doc.add_paragraph('Título do Documento')
# titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto
# run = titulo.runs[0]
# run.font.name = 'Arial'  # Define a fonte como Arial
# run.font.size = Pt(48)  # Define o tamanho da fonte como 48 pontos
# run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta

# # Adiciona mais espaço após o título (três quebras de linha)
# doc.add_paragraph('\n\n\n')

# # Adiciona uma tabela ao documento e a centraliza
# table = doc.add_table(rows=10, cols=3)
# table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza a tabela

# # Ajusta a largura da primeira coluna para 1.75 polegadas (convertido para twips) e as outras colunas para 1 polegada
# table.columns[0].width = int(8 * 1440)  # 1 polegada = 1440 twips
# table.columns[1].width = int(1 * 1440)  # 1 polegada = 1440 twips
# table.columns[2].width = int(1 * 1440)  # 1 polegada = 1440 twips

# # Dados para preencher a tabela
# dados = [
#     ("BIMESTRE", "-", "4º"),
#     ("TIPO DE PROVA", "-", "P1 OBJ"),
#     ("SÉRIE", "-", "1A"),
#     ("DISCIPLINA", "-", "MATEMÁTICA"),
#     ("PROFESSOR", "-", "THIAGO"),
#     ("Nº DE PROVAS", "-", "33"),
#     ("Nº DE 2ª CH", "-", "1"),
#     ("DATA APLICAÇÃO", "-", "09/10/23"),
#     ("DATA DIGITALIZAÇÃO", "-", "11/10/23"),
#     ("ENTREGUE POR", "-", "ESTEVAN"),
#     ("RECEBIDO POR", "-", "ALEX"),
#     ("APLICADOR", "-", "EDUARDA")
# ]

# # Define a fonte, o tamanho e converte o texto para maiúsculas na tabela
# for row_idx, row in enumerate(table.rows):
#     for col_idx, cell in enumerate(row.cells):
#         texto = dados[row_idx][col_idx].upper()  # Converte o texto para maiúsculas
#         cell.text = texto  # Preenche a célula com os dados correspondentes
#         cell.paragraphs[0].runs[0].font.name = 'Arial'  # Define a fonte como Arial
#         cell.paragraphs[0].runs[0].font.size = Pt(22)  # Define o tamanho da fonte para 22 pontos
#         cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta
        
#         # Alinha o texto à esquerda na primeira coluna
#         if col_idx == 0:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#         # Ajusta o recuo para a direita na coluna do meio
#         elif col_idx == 1:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         else:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto nas outras colunas

# # Salva o documento
# doc.save('tabela_com_titulo_e_texto_maiusculo_e_alinhamento_e_largura_da_primeira_coluna.docx')

# print('Documento criado com sucesso!')




# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# # Cria um novo documento
# doc = Document()

# # Adiciona um título ao documento com fonte Arial, tamanho 48, centralizado e cor preta (sem negrito)
# titulo = doc.add_paragraph('Título do Documento')
# titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto
# run = titulo.runs[0]
# run.font.name = 'Arial'  # Define a fonte como Arial
# run.font.size = Pt(48)  # Define o tamanho da fonte como 48 pontos
# run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta

# # Adiciona mais espaço após o título (três quebras de linha)
# doc.add_paragraph('\n\n\n')

# # Adiciona uma tabela ao documento e a centraliza
# table = doc.add_table(rows=10, cols=3)
# table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza a tabela

# # Ajusta a largura da primeira linha para 2 polegadas e as outras linhas para 1 polegada
# for col_idx, cell in enumerate(table.rows[0].cells):
#     cell.width = int(7 * 1440)  # 2 polegadas = 2880 twips

# # Dados para preencher a tabela
# dados = [
#     ("BIMESTRE", "-", "4º"),
#     ("TIPO DE PROVA", "-", "P1 OBJ"),
#     ("SÉRIE", "-", "1A"),
#     ("DISCIPLINA", "-", "MATEMÁTICA"),
#     ("PROFESSOR", "-", "THIAGO"),
#     ("Nº DE PROVAS", "-", "33"),
#     ("Nº DE 2ª CH", "-", "1"),
#     ("DATA APLICAÇÃO", "-", "09/10/23"),
#     ("DATA DIGITALIZAÇÃO", "-", "11/10/23"),
#     ("ENTREGUE POR", "-", "ESTEVAN"),
#     ("RECEBIDO POR", "-", "ALEX"),
#     ("APLICADOR", "-", "EDUARDA")
# ]

# # Define a fonte, o tamanho e converte o texto para maiúsculas na tabela
# for row_idx, row in enumerate(table.rows):
#     for col_idx, cell in enumerate(row.cells):
#         texto = dados[row_idx][col_idx].upper()  # Converte o texto para maiúsculas
#         cell.text = texto  # Preenche a célula com os dados correspondentes
#         cell.paragraphs[0].runs[0].font.name = 'Arial'  # Define a fonte como Arial
#         cell.paragraphs[0].runs[0].font.size = Pt(20)  # Define o tamanho da fonte para 22 pontos
#         cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta
        
#         # Alinha o texto à esquerda na primeira coluna
#         if col_idx == 0:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#         # Ajusta o recuo para a direita na coluna do meio
#         elif col_idx == 1:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         else:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto nas outras colunas

# # Salva o documento
# doc.save('test.docx')

# print('Documento criado com sucesso!')


# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# # Função para ajustar o texto na célula
# def ajustar_texto(texto, tamanho_fonte, largura_celula):
#     numero_caracteres = int(largura_celula / tamanho_fonte * 2)  # Fator de ajuste para a largura da célula
#     texto_ajustado = texto[:numero_caracteres]
#     return texto_ajustado

# # Cria um novo documento
# doc = Document()

# # Adiciona um título ao documento com fonte Arial, tamanho 48, centralizado e cor preta (sem negrito)
# titulo = doc.add_paragraph('Título do Documento')
# titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto
# run = titulo.runs[0]
# run.font.name = 'Arial'  # Define a fonte como Arial
# run.font.size = Pt(48)  # Define o tamanho da fonte como 48 pontos
# run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta

# # Adiciona mais espaço após o título (três quebras de linha)
# doc.add_paragraph('\n\n\n')

# # Adiciona uma tabela ao documento e a centraliza
# table = doc.add_table(rows=10, cols=3)
# table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza a tabela

# # Ajusta a largura da primeira linha para 2 polegadas e as outras linhas para 1 polegada
# largura_celula = int(2 * 1440)  # 2 polegadas = 2880 twips
# for col_idx, cell in enumerate(table.rows[0].cells):
#     cell.width = largura_celula

# # Dados para preencher a tabela
# dados = [
#     ("BIMESTRE", "-", "4º"),
#     ("TIPO DE PROVA", "-", "P1 OBJ"),
#     ("SÉRIE", "-", "1A"),
#     ("DISCIPLINA", "-", "MATEMÁTICA"),
#     ("PROFESSOR", "-", "THIAGO"),
#     ("Nº DE PROVAS", "-", "33"),
#     ("Nº DE 2ª CH", "-", "1"),
#     ("DATA APLICAÇÃO", "-", "09/10/23"),
#     ("DATA DIGITALIZAÇÃO", "-", "11/10/23"),
#     ("ENTREGUE POR", "-", "ESTEVAN"),
#     ("RECEBIDO POR", "-", "ALEX"),
#     ("APLICADOR", "-", "EDUARDA")
# ]

# # Define a fonte e o tamanho
# fonte = 'Arial'
# tamanho_fonte = 20

# # Preenche a tabela com os dados ajustados
# for row_idx, row in enumerate(table.rows):
#     for col_idx, cell in enumerate(row.cells):
#         texto = dados[row_idx][col_idx].upper()  # Converte o texto para maiúsculas
#         texto_ajustado = ajustar_texto(texto, tamanho_fonte, largura_celula)
#         cell.text = texto_ajustado  # Preenche a célula com os dados ajustados
        
#         # Configura a fonte
#         cell.paragraphs[0].runs[0].font.name = fonte
#         cell.paragraphs[0].runs[0].font.size = Pt(tamanho_fonte)
#         cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta
        
#         # Alinha o texto à esquerda na primeira coluna
#         if col_idx == 0:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#         # Ajusta o recuo para a direita na coluna do meio
#         elif col_idx == 1:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         else:
#             cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto nas outras colunas

# # Salva o documento
# doc.save('tabela_com_titulo_e_texto_maiusculo_e_alinhamento_e_largura_da_primeira_linha_ampliada.docx')

# print('Documento criado com sucesso!')



from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Cria um novo documento
doc = Document()

# Adiciona um título ao documento com fonte Arial, tamanho 48, centralizado e cor preta (sem negrito)
titulo = doc.add_paragraph('Título do Documento')
titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto
run = titulo.runs[0]
run.font.name = 'Arial'  # Define a fonte como Arial
run.font.size = Pt(48)  # Define o tamanho da fonte como 48 pontos
run.font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta

# Adiciona mais espaço após o título (três quebras de linha)
doc.add_paragraph('\n\n\n')

# Adiciona uma tabela ao documento e a centraliza
table = doc.add_table(rows=10, cols=3)
table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza a tabela

# Ajusta a largura da primeira linha para 2 polegadas e as outras linhas para 1 polegada
largura_celula = int(2 * 1440)  # 2 polegadas = 2880 twips
for col_idx, cell in enumerate(table.rows[1].cells):
    cell.width = largura_celula

# Define uma altura mínima para a primeira linha
altura_minima = int(0.5 * 1440)  # 0.5 polegadas = 720 twips
table.rows[1].height = altura_minima

# Dados para preencher a tabela
dados = [
    ("BIMESTRE", "-", "4º"),
    ("TIPO DE PROVA", "-", "P1 OBJ"),
    ("SÉRIE", "-", "1A"),
    ("DISCIPLINA", "-", "MATEMÁTICA"),
    ("PROFESSOR", "-", "THIAGO"),
    ("Nº DE PROVAS", "-", "33"),
    ("Nº DE 2ª CH", "-", "1"),
    ("DATA APLICAÇÃO", "-", "09/10/23"),
    ("DATA DIGITALIZAÇÃO", "-", "11/10/23"),
    ("ENTREGUE POR", "-", "ESTEVAN"),
    ("RECEBIDO POR", "-", "ALEX"),
    ("APLICADOR", "-", "EDUARDA")
]

# Define a fonte e o tamanho
fonte = 'Arial'
tamanho_fonte = 20

# Preenche a tabela com os dados ajustados
for row_idx, row in enumerate(table.rows):
    for col_idx, cell in enumerate(row.cells):
        texto = dados[row_idx][col_idx].upper()  # Converte o texto para maiúsculas
        cell.text = texto  # Preenche a célula com os dados correspondentes
        
        # Configura a fonte
        cell.paragraphs[0].runs[0].font.name = fonte
        cell.paragraphs[0].runs[0].font.size = Pt(tamanho_fonte)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)  # Define a cor da letra como preta
        
        # Alinha o texto à esquerda na primeira coluna
        if col_idx == 0:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # Ajusta o recuo para a direita na coluna do meio
        elif col_idx == 1:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Centraliza o texto nas outras colunas

# Salva o documento
doc.save('tabela_com_titulo_e_texto_maiusculo_e_alinhamento_e_altura_da_primeira_linha_ampliada.docx')

print('Documento criado com sucesso!')









