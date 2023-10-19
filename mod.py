# from docx import Document
# from docx.shared import Pt, RGBColor
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# # Abre o documento existente
# doc = Document('1711.1722-OBJ.docx')


# # Função para copiar o estilo de uma run para outra
# def copy_run_style(source_run, target_run):
#     target_run.bold = source_run.bold
#     target_run.italic = source_run.italic
#     target_run.underline = source_run.underline
#     target_run.font.name = source_run.font.name
#     target_run.font.size = source_run.font.size
#     target_run.font.color.rgb = source_run.font.color.rgb


# # Itera sobre as tabelas do documento
# for table in doc.tables:
#     for row in table.rows:
#         for cell in row.cells:
#             # Itera sobre os parágrafos na célula
#             for paragraph in cell.paragraphs:
#                 for run in paragraph.runs:
#                     # Substitui o texto na célula
#                     if 'bim' in run.text:
#                         run.text = run.text.replace('bim', '4')
#                         # Copia o estilo original de volta à célula após alterar o texto
#                         copy_run_style(run, paragraph.runs[0])

# # Salva as alterações no documento
# doc.save('seu_documento_modificado.docx')

# print('Texto na tabela alterado mantendo a formatação original!')


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Abre o documento existente
doc = Document('1711.1722-OBJ.docx')

# Função para copiar o estilo de uma run para outra
def copy_run_style(source_run, target_run):
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb

# Função para alterar o texto preservando o estilo
def alterar_texto_preservando_estilo(run, novo_texto):
    copy_run_style(run, run._r)
    run.text = novo_texto

# Itera sobre as tabelas do documento
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            # Itera sobre os parágrafos na célula
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    # Substitui o texto na célula preservando o estilo
                    if 'texto_antigo' in run.text:
                        alterar_texto_preservando_estilo(run, 'novo_texto')

# Para o título
titulo_paragraph = doc.paragraphs[0]  # Supondo que o título esteja no primeiro parágrafo
for run in titulo_paragraph.runs:
    if '1711' in run.text:
        alterar_texto_preservando_estilo(run, '1552 - 1553')

# Salva as alterações no documento
doc.save('seu_documento_modificado.docx')

print('Texto na tabela e no título alterado mantendo a formatação original!')
